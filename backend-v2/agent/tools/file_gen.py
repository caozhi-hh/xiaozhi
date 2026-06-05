"""文件生成工具 — 根据内容生成 Word/Excel/PDF/TXT/CSV/MD 文件供下载"""
import os
import logging
from datetime import datetime
from langchain_core.tools import tool

logger = logging.getLogger("file_gen")

# 文件存储目录
FILES_DIR = os.environ.get("FILES_DIR", "/data/files")
if not os.path.isabs(FILES_DIR):
    FILES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), FILES_DIR)

# 后端地址（用于生成下载链接）
BACKEND_URL = os.environ.get("NEXT_PUBLIC_API_URL", os.environ.get("BACKEND_URL", "https://powercz-xiaozhi.hf.space"))


def _ensure_dir():
    """确保文件目录存在"""
    os.makedirs(FILES_DIR, exist_ok=True)


def _safe_filename(name: str) -> str:
    """安全化文件名，去掉特殊字符"""
    import re
    name = re.sub(r'[\\/:*?"<>|]', '_', name)
    return name.strip()


@tool
def generate_file(filename: str, content: str, file_format: str = "txt") -> str:
    """根据内容生成文件并返回下载链接。当用户要求导出文件、生成Word文档、做表格、写文档时必须调用此工具。绝对不能假装生成了文件，必须调用此工具才能真正生成可下载的文件。

    Args:
        filename: 文件名（不含扩展名），如"AI总结报告"
        content: 文件内容（文本格式）
        file_format: 文件格式，支持 docx/xlsx/txt/csv/md/pdf

    Returns:
        下载链接或错误信息
    """
    _ensure_dir()
    filename = _safe_filename(filename)
    fmt = file_format.lower().strip()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    try:
        if fmt == "docx":
            return _gen_docx(filename, content, timestamp)
        elif fmt == "xlsx":
            return _gen_xlsx(filename, content, timestamp)
        elif fmt == "pdf":
            return _gen_pdf(filename, content, timestamp)
        elif fmt in ("txt", "csv", "md"):
            return _gen_text(filename, content, timestamp, fmt)
        else:
            # 默认 txt
            return _gen_text(filename, content, timestamp, "txt")
    except Exception as e:
        logger.error(f"文件生成失败: {e}")
        return f"文件生成失败: {e}"


def _gen_docx(filename: str, content: str, timestamp: str) -> str:
    """生成 Word 文档"""
    from docx import Document

    doc = Document()
    # 按段落分割内容
    paragraphs = content.split("\n")
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        # 简单标题检测
        if p.startswith("# "):
            doc.add_heading(p[2:], level=1)
        elif p.startswith("## "):
            doc.add_heading(p[3:], level=2)
        elif p.startswith("### "):
            doc.add_heading(p[4:], level=3)
        elif p.startswith("- ") or p.startswith("* "):
            doc.add_paragraph(p[2:], style="List Bullet")
        else:
            doc.add_paragraph(p)

    fname = f"{filename}_{timestamp}.docx"
    path = os.path.join(FILES_DIR, fname)
    doc.save(path)
    logger.info(f"生成 Word 文件: {fname}")
    return f"文件已生成！\n[点击下载: {fname}]({BACKEND_URL}/files/{fname})\n文件格式: Word (.docx)"


def _gen_xlsx(filename: str, content: str, timestamp: str) -> str:
    """生成 Excel 文件"""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = filename[:31] if len(filename) > 31 else filename

    # 解析内容为表格：支持逗号分隔或换行分隔
    lines = content.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 逗号或制表符分隔
        if "\t" in line:
            cells = line.split("\t")
        elif "," in line:
            cells = [c.strip() for c in line.split(",")]
        else:
            cells = [line]
        ws.append(cells)

    # 自动调整列宽
    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

    fname = f"{filename}_{timestamp}.xlsx"
    path = os.path.join(FILES_DIR, fname)
    wb.save(path)
    logger.info(f"生成 Excel 文件: {fname}")
    return f"文件已生成！\n[点击下载: {fname}]({BACKEND_URL}/files/{fname})\n文件格式: Excel (.xlsx)"


def _gen_pdf(filename: str, content: str, timestamp: str) -> str:
    """生成 PDF 文件（支持中文）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_LEFT

    # 注册中文字体（使用系统自带字体）
    _registered = False
    for font_path in [
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",  # Linux
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",         # Linux 备选
        "C:/Windows/Fonts/msyh.ttc",                                # Windows 微软雅黑
        "/System/Library/Fonts/PingFang.ttc",                       # macOS
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("Chinese", font_path, subfontIndex=0))
                _registered = True
                break
            except Exception:
                continue

    fname = f"{filename}_{timestamp}.pdf"
    path = os.path.join(FILES_DIR, fname)

    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm, leftMargin=20*mm, rightMargin=20*mm)

    styles = getSampleStyleSheet()
    if _registered:
        style_normal = ParagraphStyle("ChineseNormal", parent=styles["Normal"], fontName="Chinese", fontSize=11, leading=18)
        style_title = ParagraphStyle("ChineseTitle", parent=styles["Heading1"], fontName="Chinese", fontSize=16, leading=24)
        style_h2 = ParagraphStyle("ChineseH2", parent=styles["Heading2"], fontName="Chinese", fontSize=13, leading=20)
    else:
        # 没有中文字体时降级，中文会显示为方块但至少不报错
        style_normal = ParagraphStyle("Fallback", parent=styles["Normal"], fontSize=11, leading=18)
        style_title = ParagraphStyle("FallbackTitle", parent=styles["Heading1"], fontSize=16, leading=24)
        style_h2 = ParagraphStyle("FallbackH2", parent=styles["Heading2"], fontSize=13, leading=20)

    elements = []
    for line in content.split("\n"):
        stripped = line.strip()
        if not stripped:
            elements.append(Spacer(1, 4*mm))
            continue
        # 简单标题检测
        if stripped.startswith("# "):
            elements.append(Paragraph(stripped[2:].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style_title))
        elif stripped.startswith("## "):
            elements.append(Paragraph(stripped[3:].replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style_h2))
        else:
            elements.append(Paragraph(stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), style_normal))

    doc.build(elements)
    logger.info(f"生成 PDF 文件（中文支持）: {fname}")
    return f"文件已生成！\n[点击下载: {fname}]({BACKEND_URL}/files/{fname})\n文件格式: PDF (.pdf)"


def _gen_text(filename: str, content: str, timestamp: str, ext: str) -> str:
    """生成纯文本文件（txt/csv/md）"""
    fname = f"{filename}_{timestamp}.{ext}"
    path = os.path.join(FILES_DIR, fname)

    with open(path, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"生成 {ext} 文件: {fname}")
    format_names = {"txt": "文本 (.txt)", "csv": "CSV 表格 (.csv)", "md": "Markdown (.md)"}
    return f"文件已生成！\n[点击下载: {fname}]({BACKEND_URL}/files/{fname})\n文件格式: {format_names.get(ext, ext)}"
